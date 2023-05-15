import os
import pytesseract
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches


def ocr(image_folder, output_docx):
    # Create a new Word document
    doc = Document()

    # Traverse the image files in the folder
    image_files = [f for f in os.listdir(image_folder) if os.path.isfile(os.path.join(image_folder, f))]

    for image_file in image_files:
        # Build the complete path of the image file
        image_path = os.path.join(image_folder, image_file)

        # Use Tesseract-OCR for recognition
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, config=r"--psm 6 --oem 3 -l eng")

        # Create a new table row
        table = doc.add_table(rows=1, cols=3)
        row = table.rows[0]
        row.cells[0].text = 'Processed Image'
        row.cells[1].text = 'Original Image'
        row.cells[2].text = 'Text'

        # Add processed image, original image, and recognized text to the table
        row = table.add_row().cells
        cell = row[0]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        processed_image = process_image(image, text)
        processed_image_path = f"processed_{image_file}"
        processed_image.save(processed_image_path)
        run.add_picture(processed_image_path, width=Inches(2))  # Adjust the width as needed

        cell = row[1]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(2))  # Adjust the width as needed

        row[2].text = text

        print(text)

        # Delete the processed image file
        os.remove(processed_image_path)

    # Save the Word document
    doc.save(output_docx)


def process_image(image, text):
    # Add OCR marks to the image
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype("arial.ttf", 16)  # Change the font and size as needed
    draw.text((10, 10), text, font=font, fill=(255, 0, 0))  # Change the position and color as needed
    processed_image = image  # Placeholder, replace with your own image processing code
    return processed_image


# Call the OCR function
if __name__ == '__main__':
    ocr('images', 'output.docx')
